"""
文档分块处理器使用示例
演示如何使用DocumentChunker处理文档并存储到数据库
"""

import asyncio
import os
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
from document_chunker import DocumentChunker


def read_word_document(file_path: str) -> str:
    """
    从Word文档中读取文本内容
    
    Args:
        file_path: Word文档文件路径
        
    Returns:
        提取的文本内容
    """
    try:
        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 检查文件扩展名
        file_ext = Path(file_path).suffix.lower()
        if file_ext not in ['.docx', '.doc']:
            raise ValueError(f"不支持的文件格式: {file_ext}，仅支持 .docx 和 .doc 文件")
        
        # 使用改进的Word文档读取方法
        content = _extract_word_content_standalone(file_path)
        
        print(f"成功读取Word文档: {file_path}")
        print(f"总字符数: {len(content)}")
        
        return content
        
    except Exception as e:
        print(f"读取Word文档失败: {e}")
        raise


def _extract_word_content_standalone(file_path: str) -> str:
    """
    独立的Word文档内容提取函数
    
    Args:
        file_path: Word文档文件路径
        
    Returns:
        提取的完整文本内容
    """
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    
    doc = Document(file_path)
    full_text = []
    
    def extract_table_text(table):
        """提取表格中的文本"""
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        cell_text.append(paragraph.text.strip())
                if cell_text:
                    row_text.append(' '.join(cell_text))
            if row_text:
                table_text.append(' | '.join(row_text))
        return '\n'.join(table_text)
    
    # 方法1: 遍历文档中的所有元素
    try:
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # 段落
                paragraph = Paragraph(element, doc)
                text = paragraph.text.strip()
                if text:
                    full_text.append(text)
            elif isinstance(element, CT_Tbl):
                # 表格
                table = Table(element, doc)
                table_content = extract_table_text(table)
                if table_content:
                    full_text.append(table_content)
    except Exception as e:
        print(f"方法1提取失败: {e}")
    
    # 如果方法1没有提取到内容，使用备用方法
    if not full_text:
        print("使用备用方法提取Word文档内容")
        
        # 备用方法: 直接提取所有段落和表格
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)
        
        for table in doc.tables:
            table_content = extract_table_text(table)
            if table_content:
                full_text.append(table_content)
    
    # 如果还是没有内容，尝试更深层的提取
    if not full_text:
        print("使用深度提取方法")
        for paragraph in doc.paragraphs:
            paragraph_text = []
            for run in paragraph.runs:
                if run.text.strip():
                    paragraph_text.append(run.text.strip())
            if paragraph_text:
                full_text.append(' '.join(paragraph_text))
    
    result = '\n\n'.join(full_text)
    
    if not result.strip():
        print("警告: 无法从Word文档中提取任何文本内容")
        # 尝试读取文档的基本信息
        try:
            core_props = doc.core_properties
            if hasattr(core_props, 'title') and core_props.title:
                result = f"文档标题: {core_props.title}"
                print(f"仅提取到文档标题: {core_props.title}")
        except:
            pass
    
    return result


async def process_word_document(chunker: DocumentChunker, file_path: str, doc_category: str, chunk_size: int = 512):
    """
    处理Word文档并存储到数据库
    
    Args:
        chunker: 文档分块处理器实例
        file_path: Word文档文件路径
        doc_category: 文档类别
        chunk_size: 分块大小
        
    Returns:
        插入记录的ID列表
    """
    try:
        print(f"开始处理Word文档: {file_path}")
        
        # 读取Word文档内容
        text_content = read_word_document(file_path)
        
        if not text_content.strip():
            print("警告: Word文档内容为空")
            return []
        
        # 从文件路径提取文档标题
        from pathlib import Path
        doc_title = Path(file_path).stem
        
        # 使用chunker处理文本内容
        inserted_ids = await chunker.process_text_content(
            text_content=text_content,
            doc_category=doc_category,
            doc_title=doc_title,
            chunk_size=chunk_size
        )
        
        print(f"Word文档处理完成！插入了 {len(inserted_ids)} 条记录")
        print(f"记录ID: {inserted_ids}")
        
        return inserted_ids
        
    except Exception as e:
        print(f"处理Word文档失败: {e}")
        raise


async def main():
    """主函数示例"""
    # 创建文档分块处理器实例
    chunker = DocumentChunker(region_name="us-east-1")
    
    try:
        # 初始化数据库连接
        await chunker.init_db_client()
        print("数据库连接初始化成功")
        
        # 示例1: 处理文本内容
        sample_text = """
        人工智能（Artificial Intelligence，AI）是计算机科学的一个分支，它企图了解智能的实质，
        并生产出一种新的能以人类智能相似的方式做出反应的智能机器。
        
        机器学习是人工智能的核心，是使计算机具有智能的根本途径。机器学习的应用已经遍及人工智能的各个分支，
        如专家系统、自动定理证明、自然语言理解、模式识别、计算机视觉等领域。
        
        深度学习是机器学习的一个分支，它基于人工神经网络的研究，特别是利用多层次的神经网络来进行学习和表示。
        深度学习通过组合低层特征形成更加抽象的高层表示属性类别或特征，以发现数据的分布式特征表示。
        
        自然语言处理（Natural Language Processing，NLP）是人工智能和语言学领域的分支学科。
        此领域探讨如何处理及运用自然语言；自然语言处理包括多个方面和步骤，基本有认知、理解、生成等部分。
        """
        
        print("开始处理示例文本...")
        inserted_ids = await chunker.process_text_content(
            text_content=sample_text,
            doc_category="人工智能基础",
            doc_title="人工智能基础知识",
            chunk_size=64  # 较小的分块大小用于演示
        )
        
        print(f"文本处理完成！插入了 {len(inserted_ids)} 条记录")
        print(f"记录ID: {inserted_ids}")
        
        # 示例2: 处理Word文档
        # 请将下面的路径替换为你的实际Word文档路径
        word_file_path = "sample_document.docx"  # 替换为实际的Word文档路径
        
        if os.path.exists(word_file_path):
            print(f"\n开始处理Word文档: {word_file_path}")
            word_ids = await process_word_document(
                chunker=chunker,
                file_path=word_file_path,
                doc_category="Word文档",
                chunk_size=512
            )
            print(f"Word文档处理完成！插入了 {len(word_ids)} 条记录")
        else:
            print(f"\nWord文档不存在: {word_file_path}")
            print("请将 word_file_path 变量设置为实际的Word文档路径")
        
        # 示例3: 批量处理目录中的所有文档
        # documents_dir = "./documents"  # 替换为实际的文档目录路径
        # if os.path.exists(documents_dir):
        #     print(f"\n开始批量处理目录: {documents_dir}")
        #     batch_result = await process_directory_documents(
        #         directory_path=documents_dir,
        #         doc_category="批量文档",
        #         chunk_size=512,
        #         recursive=True,  # 递归处理子目录
        #         max_concurrent=3  # 最大并发数
        #     )
        #     print(f"批量处理完成！")
        #     print(f"总文件数: {batch_result['total_files']}")
        #     print(f"成功处理: {batch_result['processed_files']}")
        #     print(f"处理失败: {batch_result['failed_files']}")
        #     print(f"总记录数: {batch_result['total_records']}")
        # else:
        #     print(f"\n文档目录不存在: {documents_dir}")
        
        # 示例4: 如果有文本文件，可以这样处理
        # sample_file = "sample_document.txt"
        # if os.path.exists(sample_file):
        #     print(f"开始处理文件: {sample_file}")
        #     file_ids = await chunker.process_document_file(
        #         file_path=sample_file,
        #         doc_category="教学文档",
        #         chunk_size=512
        #     )
        #     print(f"文件处理完成！插入了 {len(file_ids)} 条记录")
        
    except Exception as e:
        print(f"处理过程中发生错误: {e}")
        
    finally:
        # 关闭数据库连接
        await chunker.close_connections()
        print("数据库连接已关闭")


async def process_word_file_standalone(file_path: str, doc_category: str = "Word文档", chunk_size: int = 512):
    """
    独立的Word文档处理函数，可以直接调用
    
    Args:
        file_path: Word文档文件路径
        doc_category: 文档类别，默认为"Word文档"
        chunk_size: 分块大小，默认为512
        
    Returns:
        插入记录的ID列表
    """
    chunker = DocumentChunker(region_name="us-east-1")
    
    try:
        await chunker.init_db_client()
        print("数据库连接初始化成功")
        
        inserted_ids = await process_word_document(
            chunker=chunker,
            file_path=file_path,
            doc_category=doc_category,
            chunk_size=chunk_size
        )
        
        return inserted_ids
        
    finally:
        await chunker.close_connections()
        print("数据库连接已关闭")


def get_supported_files(directory_path: str, recursive: bool = True) -> List[Path]:
    """
    获取目录中所有支持的文档文件
    
    Args:
        directory_path: 目录路径
        recursive: 是否递归搜索子目录
        
    Returns:
        支持的文件路径列表
    """
    supported_extensions = {'.txt', '.docx', '.doc'}
    directory = Path(directory_path)
    
    if not directory.exists():
        raise FileNotFoundError(f"目录不存在: {directory_path}")
    
    if not directory.is_dir():
        raise ValueError(f"路径不是目录: {directory_path}")
    
    files = []
    
    if recursive:
        # 递归搜索所有子目录
        for ext in supported_extensions:
            files.extend(directory.rglob(f"*{ext}"))
    else:
        # 只搜索当前目录
        for ext in supported_extensions:
            files.extend(directory.glob(f"*{ext}"))
    
    # 过滤掉隐藏文件和临时文件
    filtered_files = []
    for file_path in files:
        if not file_path.name.startswith('.') and not file_path.name.startswith('~'):
            filtered_files.append(file_path)
    
    return sorted(filtered_files)


async def process_directory_documents(
    directory_path: str,
    doc_category: str = "批量文档",
    chunk_size: int = 512,
    recursive: bool = True,
    max_concurrent: int = 1
) -> Dict[str, Any]:
    """
    批量处理目录中的所有文档
    
    Args:
        directory_path: 文档目录路径
        doc_category: 文档类别
        chunk_size: 分块大小
        recursive: 是否递归处理子目录
        max_concurrent: 最大并发处理数量
        
    Returns:
        处理结果统计信息
    """
    print(f"开始批量处理目录: {directory_path}")
    print(f"文档类别: {doc_category}")
    print(f"分块大小: {chunk_size}")
    print(f"递归处理: {recursive}")
    print(f"最大并发数: {max_concurrent}")
    print("-" * 60)
    
    # 获取所有支持的文件
    try:
        files = get_supported_files(directory_path, recursive)
        print(f"发现 {len(files)} 个支持的文档文件")
        
        if not files:
            print("没有找到支持的文档文件")
            return {
                'total_files': 0,
                'processed_files': 0,
                'failed_files': 0,
                'total_records': 0,
                'results': []
            }
        
        # 显示文件列表
        print("\n文件列表:")
        for i, file_path in enumerate(files, 1):
            print(f"  {i}. {file_path}")
        print()
        
    except Exception as e:
        print(f"获取文件列表失败: {e}")
        raise
    
    # 初始化处理器
    chunker = DocumentChunker(region_name="us-east-1")
    
    try:
        await chunker.init_db_client()
        print("数据库连接初始化成功\n")
        
        # 处理结果统计
        results = []
        total_records = 0
        processed_count = 0
        failed_count = 0
        
        # 使用信号量控制并发数量
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def process_single_file(file_path: Path) -> Dict[str, Any]:
            """处理单个文件"""
            async with semaphore:
                file_result = {
                    'file_path': str(file_path),
                    'file_name': file_path.name,
                    'success': False,
                    'records_count': 0,
                    'record_ids': [],
                    'error': None,
                    'processing_time': 0
                }
                
                start_time = asyncio.get_event_loop().time()
                
                try:
                    print(f"🔄 处理文件: {file_path.name}")
                    
                    # 使用chunker的process_document_file方法
                    inserted_ids = await chunker.process_document_file(
                        file_path=str(file_path),
                        doc_category=doc_category,
                        chunk_size=chunk_size
                    )
                    
                    file_result['success'] = True
                    file_result['records_count'] = len(inserted_ids)
                    file_result['record_ids'] = inserted_ids
                    
                    print(f"✅ 完成: {file_path.name} - 插入 {len(inserted_ids)} 条记录")
                    
                except Exception as e:
                    file_result['error'] = str(e)
                    print(f"❌ 失败: {file_path.name} - {e}")
                
                finally:
                    end_time = asyncio.get_event_loop().time()
                    file_result['processing_time'] = round(end_time - start_time, 2)
                
                return file_result
        
        # 并发处理所有文件
        print("开始并发处理文件...\n")
        tasks = [process_single_file(file_path) for file_path in files]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # 统计结果
        for result in results:
            if isinstance(result, Exception):
                failed_count += 1
                print(f"❌ 处理异常: {result}")
            elif isinstance(result, dict):
                if result['success']:
                    processed_count += 1
                    total_records += result['records_count']
                else:
                    failed_count += 1
        
        # 输出统计信息
        print("\n" + "=" * 60)
        print("📊 批量处理完成统计:")
        print(f"📁 总文件数: {len(files)}")
        print(f"✅ 成功处理: {processed_count}")
        print(f"❌ 处理失败: {failed_count}")
        print(f"📝 总记录数: {total_records}")
        print(f"⏱️  处理时间: {sum(r.get('processing_time', 0) for r in results if isinstance(r, dict)):.2f}秒")
        
        # 显示失败的文件
        if failed_count > 0:
            print(f"\n❌ 失败的文件:")
            for result in results:
                if isinstance(result, dict) and not result['success']:
                    print(f"  - {result['file_name']}: {result['error']}")
        
        return {
            'total_files': len(files),
            'processed_files': processed_count,
            'failed_files': failed_count,
            'total_records': total_records,
            'results': [r for r in results if isinstance(r, dict)]
        }
        
    finally:
        await chunker.close_connections()
        print("\n数据库连接已关闭")


if __name__ == "__main__":
    # 运行示例
    asyncio.run(main())
    
    # 或者直接处理单个Word文档（取消注释下面的代码）
    # word_path = "your_document.docx"  # 替换为实际路径
    # result = asyncio.run(process_word_file_standalone(word_path, "技术文档", 512))
    # print(f"处理结果: {result}")