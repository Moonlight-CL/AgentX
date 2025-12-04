"""
文档分块处理器
使用AWS Bedrock的Cohere模型对文档进行分块处理，并存储到edu_ref_docs表中
"""

import os
import logging
from typing import List, Dict, Any, Optional
import asyncio
import json
from pathlib import Path
import boto3
from botocore.exceptions import ClientError
from dotenv import load_dotenv
from aurora_db_client import AuroraDBClient

# 加载环境变量
load_dotenv()

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentChunker:
    """文档分块处理器类"""
    
    def __init__(self, region_name: str = None):
        """
        初始化文档分块处理器
        
        Args:
            region_name: AWS区域名称，如果为None则从环境变量获取
        """
        # 从环境变量获取AWS区域，如果没有设置则使用默认值
        self.region_name = region_name or os.getenv('AWS_REGION', 'us-east-1')
        self.bedrock_client = None
        self.db_client = None
        self._init_bedrock_client()
    
    def _init_bedrock_client(self):
        """初始化Bedrock客户端"""
        try:
            self.bedrock_client = boto3.client(
                'bedrock-runtime',
                region_name=self.region_name,
                aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
                aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
                aws_session_token=os.getenv('AWS_SESSION_TOKEN')  # 可选，用于临时凭证
            )
            logger.info(f"Bedrock客户端初始化成功，区域: {self.region_name}")
        except Exception as e:
            logger.error(f"Bedrock客户端初始化失败: {e}")
            raise
    
    async def init_db_client(self):
        """初始化数据库客户端"""
        try:
            self.db_client = AuroraDBClient()
            await self.db_client.init_connection_pool()
            logger.info("数据库客户端初始化成功")
        except Exception as e:
            logger.error(f"数据库客户端初始化失败: {e}")
            raise
    
    async def close_connections(self):
        """关闭所有连接"""
        if self.db_client:
            await self.db_client.close_pool()
    
    def chunk_document_with_overlap(self, 
                                  document_text: str, 
                                  chunk_size: int = 512,
                                  overlap_percentage: float = 0.1) -> List[str]:
        """
        对文档进行分块处理，支持重叠
        
        Args:
            document_text: 文档文本内容
            chunk_size: 分块大小（字符数）
            overlap_percentage: 重叠百分比（0.0-1.0）
            
        Returns:
            分块后的文本列表
        """
        try:
            # 使用改进的文本分块方法
            chunks = self._advanced_text_chunking(document_text, chunk_size, overlap_percentage)
            logger.info(f"文档分块完成，共生成 {len(chunks)} 个分块")
            return chunks
            
        except Exception as e:
            logger.error(f"文档分块处理失败: {e}")
            raise
    
    def _simple_text_chunking(self, text: str, chunk_size: int = 512) -> List[str]:
        """
        简单的文本分块方法（备选方案，无重叠）
        
        Args:
            text: 输入文本
            chunk_size: 每个分块的大概字符数
            
        Returns:
            分块后的文本列表
        """
        words = text.split()
        chunks = []
        current_chunk = []
        current_length = 0
        
        for word in words:
            word_length = len(word) + 1  # +1 for space
            if current_length + word_length > chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = [word]
                current_length = word_length
            else:
                current_chunk.append(word)
                current_length += word_length
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        logger.info(f"使用简单分块方法，共生成 {len(chunks)} 个分块")
        return chunks
    
    def _advanced_text_chunking(self, text: str, chunk_size: int = 512, overlap_percentage: float = 0.1) -> List[str]:
        """
        高级文本分块方法，支持重叠和智能分割
        
        Args:
            text: 输入文本
            chunk_size: 每个分块的目标字符数
            overlap_percentage: 重叠百分比（0.0-1.0）
            
        Returns:
            分块后的文本列表
        """
        if not text.strip():
            return []
        
        # 计算重叠大小
        overlap_size = int(chunk_size * overlap_percentage)
        
        # 首先按段落分割
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            # 如果当前段落本身就很长，需要进一步分割
            if len(paragraph) > chunk_size:
                # 先处理当前积累的内容
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                
                # 分割长段落
                long_paragraph_chunks = self._split_long_paragraph(paragraph, chunk_size, overlap_size)
                chunks.extend(long_paragraph_chunks)
            else:
                # 检查添加这个段落是否会超过分块大小
                if len(current_chunk) + len(paragraph) + 2 > chunk_size:  # +2 for \n\n
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                        
                        # 如果需要重叠，从当前chunk的末尾取一部分作为下一个chunk的开始
                        if overlap_size > 0 and len(current_chunk) > overlap_size:
                            overlap_text = current_chunk[-overlap_size:].strip()
                            # 找到完整的词边界
                            words = overlap_text.split()
                            if len(words) > 1:
                                overlap_text = ' '.join(words[1:])  # 去掉可能被截断的第一个词
                            current_chunk = overlap_text + "\n\n" + paragraph if overlap_text else paragraph
                        else:
                            current_chunk = paragraph
                    else:
                        current_chunk = paragraph
                else:
                    if current_chunk:
                        current_chunk += "\n\n" + paragraph
                    else:
                        current_chunk = paragraph
        
        # 处理最后的chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # 过滤掉太短的chunks
        min_chunk_size = max(50, chunk_size // 10)  # 最小chunk大小为目标大小的1/10，但不少于50字符
        filtered_chunks = [chunk for chunk in chunks if len(chunk) >= min_chunk_size]
        
        logger.info(f"使用高级分块方法，共生成 {len(filtered_chunks)} 个分块（原始: {len(chunks)}）")
        return filtered_chunks
    
    def _split_long_paragraph(self, paragraph: str, chunk_size: int, overlap_size: int) -> List[str]:
        """
        分割长段落
        
        Args:
            paragraph: 长段落文本
            chunk_size: 分块大小
            overlap_size: 重叠大小
            
        Returns:
            分割后的文本块列表
        """
        chunks = []
        
        # 尝试按句子分割
        sentences = self._split_into_sentences(paragraph)
        
        current_chunk = ""
        for sentence in sentences:
            if len(current_chunk) + len(sentence) + 1 > chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    
                    # 添加重叠
                    if overlap_size > 0 and len(current_chunk) > overlap_size:
                        overlap_text = current_chunk[-overlap_size:].strip()
                        words = overlap_text.split()
                        if len(words) > 1:
                            overlap_text = ' '.join(words[1:])
                        current_chunk = overlap_text + " " + sentence if overlap_text else sentence
                    else:
                        current_chunk = sentence
                else:
                    # 如果单个句子就超过了chunk_size，按词分割
                    if len(sentence) > chunk_size:
                        word_chunks = self._split_by_words(sentence, chunk_size, overlap_size)
                        chunks.extend(word_chunks)
                        current_chunk = ""
                    else:
                        current_chunk = sentence
            else:
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """
        将文本分割成句子
        
        Args:
            text: 输入文本
            
        Returns:
            句子列表
        """
        # 简单的句子分割，基于常见的句子结束符
        import re
        
        # 中英文句子结束符
        sentence_endings = r'[.!?。！？；;]'
        sentences = re.split(sentence_endings, text)
        
        # 过滤空句子并保留标点
        result = []
        for i, sentence in enumerate(sentences[:-1]):  # 最后一个通常是空的
            sentence = sentence.strip()
            if sentence:
                result.append(sentence)
        
        # 处理最后一个句子（如果有内容）
        if sentences[-1].strip():
            result.append(sentences[-1].strip())
        
        return result
    
    def _split_by_words(self, text: str, chunk_size: int, overlap_size: int) -> List[str]:
        """
        按词分割文本
        
        Args:
            text: 输入文本
            chunk_size: 分块大小
            overlap_size: 重叠大小
            
        Returns:
            分割后的文本块列表
        """
        words = text.split()
        chunks = []
        current_chunk = []
        current_length = 0
        
        for word in words:
            word_length = len(word) + 1  # +1 for space
            
            if current_length + word_length > chunk_size and current_chunk:
                chunk_text = ' '.join(current_chunk)
                chunks.append(chunk_text)
                
                # 计算重叠的词数
                if overlap_size > 0:
                    overlap_words = []
                    overlap_length = 0
                    
                    # 从后往前取词，直到达到重叠大小
                    for i in range(len(current_chunk) - 1, -1, -1):
                        word_len = len(current_chunk[i]) + 1
                        if overlap_length + word_len <= overlap_size:
                            overlap_words.insert(0, current_chunk[i])
                            overlap_length += word_len
                        else:
                            break
                    
                    current_chunk = overlap_words + [word]
                    current_length = sum(len(w) + 1 for w in current_chunk)
                else:
                    current_chunk = [word]
                    current_length = word_length
            else:
                current_chunk.append(word)
                current_length += word_length
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    def _extract_word_content(self, file_path: str) -> str:
        """
        从Word文档中提取所有文本内容
        
        Args:
            file_path: Word文档文件路径
            
        Returns:
            提取的完整文本内容
        """
        from docx import Document
        from docx.document import Document as DocumentType
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import _Cell, Table
        from docx.text.paragraph import Paragraph
        
        doc = Document(file_path)
        full_text = []
        
        def extract_text_from_element(element):
            """递归提取元素中的文本"""
            if hasattr(element, 'text') and element.text.strip():
                return element.text.strip()
            return ""
        
        def extract_table_text(table):
            """提取表格中的文本"""
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = []
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            cell_text.append(paragraph.text.strip())
                    if cell_text:
                        row_text.append(' '.join(cell_text))
                if row_text:
                    table_text.append(' | '.join(row_text))
            return '\n'.join(table_text)
        
        # 遍历文档中的所有元素
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # 段落
                paragraph = Paragraph(element, doc)
                text = paragraph.text.strip()
                if text:
                    full_text.append(text)
            elif isinstance(element, CT_Tbl):
                # 表格
                table = Table(element, doc)
                table_content = extract_table_text(table)
                if table_content:
                    full_text.append(table_content)
        
        # 如果上面的方法没有提取到内容，使用备用方法
        if not full_text:
            logger.warning("使用标准方法未提取到内容，尝试备用方法")
            
            # 备用方法1: 直接提取所有段落
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    full_text.append(text)
            
            # 备用方法2: 提取表格内容
            for table in doc.tables:
                table_content = extract_table_text(table)
                if table_content:
                    full_text.append(table_content)
        
        # 如果还是没有内容，尝试更深层的提取
        if not full_text:
            logger.warning("使用备用方法仍未提取到内容，尝试深度提取")
            
            # 深度提取：遍历所有runs
            for paragraph in doc.paragraphs:
                paragraph_text = []
                for run in paragraph.runs:
                    if run.text.strip():
                        paragraph_text.append(run.text.strip())
                if paragraph_text:
                    full_text.append(' '.join(paragraph_text))
        
        result = '\n\n'.join(full_text)
        
        if not result.strip():
            logger.error(f"无法从Word文档中提取任何文本内容: {file_path}")
            # 尝试读取文档的基本信息
            try:
                core_props = doc.core_properties
                if hasattr(core_props, 'title') and core_props.title:
                    result = f"文档标题: {core_props.title}"
                    logger.info(f"仅提取到文档标题: {core_props.title}")
            except:
                pass
        
        logger.info(f"从Word文档提取文本，总长度: {len(result)} 字符")
        
        # 调试信息
        if len(result) < 100:
            logger.warning(f"提取的内容可能不完整，内容预览: {result[:200]}")
        
        return result
    
    def generate_embeddings_with_cohere(self, texts: List[str]) -> List[List[float]]:
        """
        使用Cohere模型生成文本嵌入向量
        
        Args:
            texts: 文本列表
            
        Returns:
            嵌入向量列表
        """
        try:
            embeddings = []
            
            # 批量处理文本，避免单次请求过大
            batch_size = 10
            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i + batch_size]
                
                request_body = {
                    "texts": batch_texts,
                    "input_type": "search_document"
                }
                
                response = self.bedrock_client.invoke_model(
                    modelId="cohere.embed-multilingual-v3",
                    contentType="application/json",
                    accept="application/json",
                    body=json.dumps(request_body)
                )
                
                response_body = json.loads(response['body'].read())
                
                if 'embeddings' in response_body:
                    embeddings.extend(response_body['embeddings'])
            
            logger.info(f"成功生成 {len(embeddings)} 个嵌入向量")
            return embeddings
            
        except Exception as e:
            logger.error(f"生成嵌入向量失败: {e}")
            # 返回零向量作为备选方案
            return [[0.0] * 1024 for _ in texts]  # Cohere默认1024维
    
    async def process_document_file(self, 
                                  file_path: str, 
                                  doc_category: str,
                                  doc_title: Optional[str] = None,
                                  chunk_size: int = 512) -> List[int]:
        """
        处理文档文件并存储到数据库
        
        Args:
            file_path: 文档文件路径
            doc_category: 文档类别
            doc_title: 文档标题，如果为None则使用文件名
            chunk_size: 分块大小
            
        Returns:
            插入记录的ID列表
        """
        try:
            # 读取文档内容
            file_path_obj = Path(file_path)
            if not file_path_obj.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            # 根据文件类型读取内容
            file_ext = file_path_obj.suffix.lower()
            if file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    document_text = f.read()
            elif file_ext in ['.docx', '.doc']:
                try:
                    from docx import Document
                    document_text = self._extract_word_content(file_path)
                except ImportError:
                    raise ValueError("处理Word文档需要安装python-docx库: pip install python-docx")
            else:
                raise ValueError(f"不支持的文件类型: {file_ext}，支持的格式: .txt, .docx, .doc")
            
            # 如果没有提供标题，使用文件名（不含扩展名）
            if doc_title is None:
                doc_title = file_path_obj.stem
            
            logger.info(f"开始处理文档: {file_path}，标题: {doc_title}")
            logger.info(f"文档内容长度: {len(document_text)} 字符")
            
            # 对文档进行分块
            chunks = self.chunk_document_with_overlap(document_text, chunk_size)
            
            if not chunks:
                logger.warning("文档分块结果为空")
                return []
            
            # 生成嵌入向量
            embeddings = self.generate_embeddings_with_cohere(chunks)
            
            # 确保数据库客户端已初始化
            if not self.db_client:
                await self.init_db_client()
            
            # 存储到数据库
            inserted_ids = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                try:
                    # 生成简单摘要（取前100个字符）
                    summary = chunk[:100] + "..." if len(chunk) > 100 else chunk
                    
                    # 准备插入数据
                    record_data = {
                        'doc_category': doc_category,
                        'doc_title': doc_title,
                        'content': chunk,
                        'content_vector': embedding,
                        'summary': summary
                    }
                    
                    # 插入记录
                    record_id = await self.db_client.insert_record('edu_ref_docs', record_data)
                    inserted_ids.append(record_id)
                    
                    logger.info(f"成功插入分块 {i+1}/{len(chunks)}，记录ID: {record_id}")
                    
                except Exception as e:
                    logger.error(f"插入分块 {i+1} 失败: {e}")
                    continue
            
            logger.info(f"文档处理完成，共插入 {len(inserted_ids)} 条记录")
            return inserted_ids
            
        except Exception as e:
            logger.error(f"处理文档文件失败: {e}")
            raise
    
    async def process_text_content(self, 
                                 text_content: str, 
                                 doc_category: str,
                                 doc_title: str,
                                 chunk_size: int = 512) -> List[int]:
        """
        处理文本内容并存储到数据库
        
        Args:
            text_content: 文本内容
            doc_category: 文档类别
            doc_title: 文档标题
            chunk_size: 分块大小
            
        Returns:
            插入记录的ID列表
        """
        try:
            logger.info("开始处理文本内容")
            
            # 对文本进行分块
            chunks = self.chunk_document_with_overlap(text_content, chunk_size)
            
            if not chunks:
                logger.warning("文本分块结果为空")
                return []
            
            # 生成嵌入向量
            embeddings = self.generate_embeddings_with_cohere(chunks)
            
            # 确保数据库客户端已初始化
            if not self.db_client:
                await self.init_db_client()
            
            # 存储到数据库
            inserted_ids = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                try:
                    # 生成简单摘要
                    summary = chunk[:100] + "..." if len(chunk) > 100 else chunk
                    
                    # 准备插入数据
                    record_data = {
                        'doc_category': doc_category,
                        'doc_title': doc_title,
                        'content': chunk,
                        'content_vector': embedding,
                        'summary': summary
                    }
                    
                    # 插入记录
                    record_id = await self.db_client.insert_record('edu_ref_docs', record_data)
                    inserted_ids.append(record_id)
                    
                    logger.info(f"成功插入分块 {i+1}/{len(chunks)}，记录ID: {record_id}")
                    
                except Exception as e:
                    logger.error(f"插入分块 {i+1} 失败: {e}")
                    continue
            
            logger.info(f"文本处理完成，共插入 {len(inserted_ids)} 条记录")
            return inserted_ids
            
        except Exception as e:
            logger.error(f"处理文本内容失败: {e}")
            raise


# 使用示例
async def example_usage():
    """使用示例"""
    chunker = DocumentChunker()
    
    try:
        # 初始化数据库连接
        await chunker.init_db_client()
        
        # 示例1: 处理文本文件
        # file_path = "sample_document.txt"
        # inserted_ids = await chunker.process_document_file(file_path, "教学材料", "示例文档标题")
        # print(f"处理文件完成，插入记录ID: {inserted_ids}")
        
        # 示例2: 处理文本内容
        sample_text = """
        这是一个示例教育文档。它包含了关于机器学习的基础知识。
        机器学习是人工智能的一个重要分支，它使计算机能够在没有明确编程的情况下学习。
        机器学习算法通过分析数据来识别模式，并使用这些模式来做出预测或决策。
        常见的机器学习类型包括监督学习、无监督学习和强化学习。
        """
        
        inserted_ids = await chunker.process_text_content(sample_text, "机器学习基础", "机器学习入门教程")
        print(f"处理文本完成，插入记录ID: {inserted_ids}")
        
    finally:
        # 关闭连接
        await chunker.close_connections()


if __name__ == "__main__":
    asyncio.run(example_usage())